from docx import Document

def create_personalized_documents(template_path, people_list):
    for person in people_list:
        # Carregar o documento template
        doc = Document(template_path)

        # Substituir os placeholders
        for paragraph in doc.paragraphs:
            if 'NOME' in paragraph.text:
                paragraph.text = paragraph.text.replace('NOME', person['nome'])
            if 'ENDERECO' in paragraph.text:
                paragraph.text = paragraph.text.replace('ENDERECO', person['endereco'])
            if 'CARGO' in paragraph.text:
                paragraph.text = paragraph.text.replace('CARGO', person['cargo'])

        # Salvar o documento personalizado
        personalized_path = f"personalized_{person['nome']}.docx"
        doc.save(personalized_path)

# Exemplo de uso
template_path = 'caminho_para_o_template.docx'
people_list = [
    {'nome': 'João', 'endereco': 'Rua Exemplo 123', 'cargo': 'Gerente'},
    {'nome': 'Maria', 'endereco': 'Avenida Exemplo 456', 'cargo': 'Diretora'},
    # ... mais pessoas
]

create_personalized_documents(template_path, people_list)


from docx import Document

def replace_placeholder(doc, placeholder, real_value):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, real_value)

doc = Document('path_to_template.docx')
replace_placeholder(doc, '[NOME]', 'João Silva')
replace_placeholder(doc, '[ENDERECO]', 'Rua Exemplo, 123')
replace_placeholder(doc, '[CARGO]', 'Gerente')
doc.save('path_to_new_document.docx')


import openpyxl
from docx import Document

# Carregar a planilha do Excel
excel_path = 'caminho_para_a_planilha.xlsx'
workbook = openpyxl.load_workbook(excel_path)
sheet = workbook.active

# Carregar o documento do Word
word_path = 'caminho_para_o_documento.docx'
doc = Document(word_path)

# Adicionar informações ao final do documento
for row in sheet.iter_rows(min_row=2, values_only=True):  # Ajuste 'min_row' conforme necessário
    nome, cargo, endereco = row
    paragraph = doc.add_paragraph()
    paragraph.add_run(f'Nome: {nome}\nCargo: {cargo}\nEndereço: {endereco}\n\n')

# Salvar o documento atualizado
doc.save('documento_atualizado.docx')


import openpyxl
from docx import Document

def find_placeholder_paragraph(document, placeholder):
    for paragraph in document.paragraphs:
        if placeholder in paragraph.text:
            return paragraph
    return None

# Carregar a planilha do Excel
excel_path = 'caminho_para_a_planilha.xlsx'
workbook = openpyxl.load_workbook(excel_path)
sheet = workbook.active

# Carregar o documento do Word
word_path = 'caminho_para_o_documento.docx'
doc = Document(word_path)

# Encontrar o placeholder no documento
placeholder_paragraph = find_placeholder_paragraph(doc, '[[INSERIR_TABELA_AQUI]]')

if placeholder_paragraph:
    # Criar uma tabela
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nome'
    hdr_cells[1].text = 'Cargo'
    hdr_cells[2].text = 'Endereço'

    # Adicionar linhas à tabela para cada pessoa
    for row in sheet.iter_rows(min_row=2, values_only=True):
        nome, cargo, endereco = row
        row_cells = table.add_row().cells
        row_cells[0].text = nome
        row_cells[1].text = cargo
        row_cells[2].text = endereco

    # Adicionar a tabela no lugar do placeholder
    p = placeholder_paragraph._element
    p.addnext(table._element)
    p.getparent().remove(p)

# Salvar o documento atualizado
doc.save('documento_atualizado_com_tabela.docx')
